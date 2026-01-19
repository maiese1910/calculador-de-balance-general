from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_word():
    doc = Document()
    
    # --- PORTADA ---
    title = doc.add_heading('UNIVERSIDAD SANTA MARÍA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n' * 5)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PROYECTO: CALCULADOR DE BALANCE GENERAL')
    run.bold = True
    run.font.size = Pt(24)
    
    doc.add_paragraph('\n' * 2)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('FACULTAD DE CIENCIAS SOCIALES Y ECONÓMICAS')
    run2.font.size = Pt(14)
    
    doc.add_page_break()
    
    # --- 1. ENTREVISTA ---
    doc.add_heading('1. Documentación de Entrevista', level=1)
    doc.add_paragraph('Contexto: Entrevista realizada a estudiantes de la Facultad de Ciencias Sociales y Económicas.')
    
    p = doc.add_paragraph()
    p.add_run('P: ¿Cuál es el mayor desafío al realizar prácticas de contabilidad?').bold = True
    doc.add_paragraph('R: El tiempo que toma cuadrar el Balance de Comprobación y el Libro Mayor. Un pequeño error al transcribir una cifra puede tomarnos horas de revisión manual.')
    
    p = doc.add_paragraph()
    p.add_run('P: ¿Cómo ayudaría una herramienta digital?').bold = True
    doc.add_paragraph('R: Permitiría validar instantáneamente si nuestros cálculos son correctos antes de entregar la tarea, sirviendo como un tutor de validación 24/7.')

    # --- 2. DIAGRAMA DE PROCESOS ---
    doc.add_heading('2. Diagrama de Procesos', level=1)
    doc.add_paragraph('Flujo de funcionamiento del sistema:')
    processes = [
        "Inicio / Validación de Licencia",
        "Carga de Libro Diario y Libro Mayor (XLSX)",
        "Procesamiento Smart-Parsing (Detección de encabezados)",
        "Clasificación automática REAL (Balance) / NOMINAL (Resultados)",
        "Cálculo de Saldos y Comprobación de Partida Doble",
        "Presentación de Resultados (Verde/Naranja)",
        "Exportación de Reportes Profesionales"
    ]
    for proc in processes:
        doc.add_paragraph(proc, style='List Bullet')

    # --- 3. DICCIONARIO DE VARIABLES ---
    doc.add_heading('3. Diccionario de Variables', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Tipo'
    hdr_cells[2].text = 'Descripción'
    
    variables = [
        ("diaryRows", "Array[Obj]", "Filas normalizadas del Libro Diario."),
        ("ledgerRows", "Array[Obj]", "Filas normalizadas del Libro Mayor."),
        ("trialBalance", "Array[Obj]", "Saldos procesados para el Balance."),
        ("type", "String", "REAL (Balance) / NOMINAL (Resultados)."),
        ("incomeStatement", "Object", "Totales de ingresos, gastos y utilidad."),
        ("isBalanced", "Boolean", "Determina si Total Debe = Total Haber.")
    ]
    
    for var, vtype, desc in variables:
        row_cells = table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = vtype
        row_cells[2].text = desc

    # --- 4. OBJETIVOS Y ALCANCE ---
    doc.add_heading('4. Objetivos y Alcance', level=1)
    
    doc.add_heading('Planteamiento del Problema', level=2)
    doc.add_paragraph('La enseñanza tradicional requiere una validación manual exhaustiva, lo que genera frustración en el estudiante y demora el aprendizaje de los conceptos analíticos.')
    
    doc.add_heading('Objetivo General', level=2)
    p = doc.add_paragraph()
    p.add_run('Optimizar la validación de actividades contables estudiantiles mediante la automatización digital de balances financieros básicos.').bold = True
    
    doc.add_heading('Objetivos Específicos', level=2)
    specifics = [
        "Implementar un motor de lectura de archivos Excel flexible.",
        "Clasificar automáticamente cuentas reales y nominales.",
        "Proporcionar feedback visual inmediato sobre el principio de partida doble."
    ]
    for s in specifics:
        doc.add_paragraph(s, style='List Bullet')
        
    doc.add_heading('Alcance', level=2)
    doc.add_paragraph('El proyecto abarca desde la lectura de libros contables hasta la generación de Estados Financieros básicos, limitado a propósitos educativos universitarios.')
    
    doc.add_heading('Ciclo de Vida', level=2)
    cycles = ["Relevamiento", "Diseño", "Implementación", "Despliegue y Retroalimentación"]
    for c in cycles:
        doc.add_paragraph(c, style='List Bullet')

    # --- 5. MANUAL DE USUARIO ---
    doc.add_heading('5. Manual de Usuario (Fines Educativos)', level=1)
    doc.add_paragraph('Instrucciones para validar actividades de la USM:')
    steps = [
        "Preparación: Asegúrate de que tu tarea de Libro Diario esté en un archivo Excel.",
        "Validación Inicial: Carga tu archivo. El sistema validará la partida doble inmediatamente.",
        "Análisis: Revisa la Clasificación para entender el destino de cada cuenta.",
        "Corrección: Si ves un mensaje naranja, verifica tus sumas y vuelve a cargar.",
        "Entrega Final: Una vez logres el balance cuadrado, exporta tu trabajo profesionalmente."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # Guardar
    output_path = 'Guia_Informativa_Proyecto.docx'
    doc.save(output_path)
    print(f"Documento '{output_path}' generado exitosamente en {os.getcwd()}")

if __name__ == "__main__":
    generate_word()
